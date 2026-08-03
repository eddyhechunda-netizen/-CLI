#!/usr/bin/env python3
"""
extract_requirements.py — 从需求文档抽取纯文本，供分析、拆解测试用例。

飞书云文档/Wiki 通过 lark-cli docs +fetch 读取为 Markdown。PDF 优先用
pdftotext -layout（保留表格的视觉排版，需求文档里大量"分类|Item|需求"
三栏表格靠空格对齐，layout 模式能最大程度保住这种结构）。无 pdftotext 时回退
pdfplumber。非 PDF（docx/txt/md）按类型读取。

目的：把"读文档"这步标准化，让 Claude 拿到干净文本就能专注做用例设计，而不必
每次纠结用什么库、怎么处理中文换行。

用法：
    python3 extract_requirements.py 需求.pdf             # 打印到 stdout
    python3 extract_requirements.py 需求.pdf -o req.txt  # 写文件
    python3 extract_requirements.py '<飞书文档URL>' -o req.md
"""
import argparse
import json
import os
import shutil
import subprocess
import sys
from pathlib import Path
from urllib.parse import urlparse


LARK_DOC_HOST_SUFFIXES = (
    ".feishu.cn",
    ".larksuite.com",
    ".doubao.com",
)


def is_lark_doc_url(value: str) -> bool:
    parsed = urlparse(value)
    return (
        parsed.scheme in {"http", "https"}
        and any(parsed.netloc.lower().endswith(suffix) for suffix in LARK_DOC_HOST_SUFFIXES)
        and any(part in parsed.path for part in ("/docx/", "/doc/", "/wiki/"))
    )


def from_lark_doc(url: str) -> str:
    if not shutil.which("lark-cli"):
        sys.exit("读取飞书文档需要 lark-cli，请先安装并完成配置。")

    env = os.environ.copy()
    env["LARKSUITE_CLI_NO_UPDATE_NOTIFIER"] = "1"
    env["LARKSUITE_CLI_NO_SKILLS_NOTIFIER"] = "1"
    out = subprocess.run(
        [
            "lark-cli", "docs", "+fetch",
            "--doc", url,
            "--doc-format", "markdown",
            "--detail", "simple",
            "--as", "user",
            "--format", "json",
        ],
        capture_output=True,
        text=True,
        timeout=120,
        env=env,
    )
    if out.returncode != 0:
        message = out.stderr.strip() or out.stdout.strip() or "未知错误"
        sys.exit(f"读取飞书文档失败：{message}")

    try:
        payload = json.loads(out.stdout)
    except json.JSONDecodeError as exc:
        sys.exit(f"lark-cli 返回的不是有效 JSON：{exc}")

    if payload.get("ok") is not True:
        error = payload.get("error") or {}
        sys.exit(f"读取飞书文档失败：{error.get('message') or payload}")

    content = (
        payload.get("data", {})
        .get("document", {})
        .get("content", "")
    )
    if not str(content).strip():
        sys.exit("飞书文档读取成功，但正文为空。")
    return str(content)


def from_pdf(path: str) -> str:
    # 1) pdftotext -layout：poppler 工具，对栏位对齐的表格保真最好
    if shutil.which("pdftotext"):
        try:
            out = subprocess.run(
                ["pdftotext", "-layout", "-enc", "UTF-8", path, "-"],
                capture_output=True, text=True, timeout=120,
            )
            if out.returncode == 0 and out.stdout.strip():
                return out.stdout
        except Exception as e:  # noqa: BLE001
            print(f"[warn] pdftotext 失败，回退 pdfplumber：{e}", file=sys.stderr)

    # 2) pdfplumber 回退
    try:
        import pdfplumber
    except ImportError:
        sys.exit("缺少 pdftotext 和 pdfplumber，至少装一个："
                 "apt install poppler-utils 或 pip install pdfplumber")
    chunks = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages, 1):
            chunks.append(f"\n===== 第 {i} 页 =====\n")
            chunks.append(page.extract_text() or "")
            # 额外把识别到的表格以管道分隔补一份，避免漏读
            for tbl in page.extract_tables():
                chunks.append("\n[表格]\n")
                for row in tbl:
                    chunks.append(" | ".join(c or "" for c in row) + "\n")
    return "".join(chunks)


def from_docx(path: str) -> str:
    try:
        import docx  # python-docx
    except ImportError:
        sys.exit("读取 .docx 需要 python-docx：pip install python-docx")
    d = docx.Document(path)
    parts = [p.text for p in d.paragraphs]
    for t in d.tables:  # 表格按行拼接
        for row in t.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


def extract(path: str) -> str:
    if is_lark_doc_url(path):
        return from_lark_doc(path)
    ext = Path(path).suffix.lower()
    if ext == ".pdf":
        return from_pdf(path)
    if ext in (".docx",):
        return from_docx(path)
    if ext in (".txt", ".md", ".markdown", ".text"):
        return Path(path).read_text(encoding="utf-8", errors="replace")
    sys.exit(f"不支持的文件类型：{ext}（支持 pdf/docx/txt/md）")


def main():
    ap = argparse.ArgumentParser(description="从需求文档抽取纯文本")
    ap.add_argument("doc", help="飞书文档/Wiki URL，或本地需求文档路径（pdf/docx/txt/md）")
    ap.add_argument("-o", "--output", help="输出文本路径；缺省打印到 stdout")
    args = ap.parse_args()

    if not is_lark_doc_url(args.doc) and not Path(args.doc).exists():
        sys.exit(f"文件不存在：{args.doc}")

    text = extract(args.doc)
    if args.output:
        Path(args.output).write_text(text, encoding="utf-8")
        nlines = text.count("\n") + 1
        print(f"✓ 已抽取 {nlines} 行文本 → {args.output}")
    else:
        sys.stdout.write(text)


if __name__ == "__main__":
    main()
